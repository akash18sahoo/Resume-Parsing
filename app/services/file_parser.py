import os
import io
from fastapi import HTTPException, status
import fitz  # PyMuPDF
from docx import Document
from PIL import Image

# Supported Extensions and mapping to unified types
# pdf, docx, text (for txt), image (for png, jpg, jpeg, webp)
SUPPORTED_EXTENSIONS = {
    ".pdf": "pdf",
    ".docx": "docx",
    ".txt": "text",
    ".png": "image",
    ".jpg": "image",
    ".jpeg": "image",
    ".webp": "image"
}

# Max file size: default 10MB, loaded dynamically from env or default
MAX_FILE_SIZE_MB = int(os.getenv("MAX_FILE_SIZE_MB", 10))
MAX_FILE_SIZE_BYTES = MAX_FILE_SIZE_MB * 1024 * 1024

def validate_file(filename: str, file_size: int) -> str:
    """
    Validates file extension and size.
    Returns the file_type string ('pdf', 'docx', 'text', 'image').
    Raises HTTPException if file is invalid.
    """
    if not filename:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Filename is missing or empty"
        )
    
    _, ext = os.path.splitext(filename.lower())
    
    if ext not in SUPPORTED_EXTENSIONS:
        supported_list = ", ".join(SUPPORTED_EXTENSIONS.keys())
        raise HTTPException(
            status_code=status.HTTP_415_UNSUPPORTED_MEDIA_TYPE,
            detail=f"Unsupported file format '{ext}'. Supported formats: {supported_list}"
        )
    
    if file_size > MAX_FILE_SIZE_BYTES:
        raise HTTPException(
            status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
            detail=f"File size exceeds maximum limit of {MAX_FILE_SIZE_MB}MB"
        )
    
    if file_size == 0:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="File is empty"
        )

    return SUPPORTED_EXTENSIONS[ext]

def extract_pdf_text(file_bytes: bytes) -> str:
    """
    Extracts text from PDF bytes using PyMuPDF (fitz).
    """
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        text_parts = []
        for page in doc:
            text_parts.append(page.get_text())
        
        extracted_text = "\n".join(text_parts).strip()
        
        if not extracted_text:
            # Try to fallback or check if scanned PDF
            raise ValueError("No text extracted from PDF. It might be scanned/image-only.")
            
        return extracted_text
    except Exception as e:
        # If it fails, or it is scanned, raise error to be handled upstream
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Failed to parse PDF text: {str(e)}. If this is a scanned PDF, please convert to image first or upload as an image."
        )

def extract_docx_text(file_bytes: bytes) -> str:
    """
    Extracts text from DOCX bytes using python-docx.
    Iterates paragraphs and tables.
    """
    try:
        doc = Document(io.BytesIO(file_bytes))
        text_parts = []
        
        # Extract paragraph text
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
                
        # Extract table text
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_parts.append(" | ".join(row_text))
                    
        extracted_text = "\n".join(text_parts).strip()
        
        if not extracted_text:
            raise ValueError("No text extracted from DOCX.")
            
        return extracted_text
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Failed to parse DOCX text: {str(e)}"
        )

def extract_txt_text(file_bytes: bytes) -> str:
    """
    Decodes plain text bytes to string. Handles common encodings.
    """
    for encoding in ("utf-8", "latin-1", "utf-16"):
        try:
            return file_bytes.decode(encoding).strip()
        except UnicodeDecodeError:
            continue
    raise HTTPException(
        status_code=status.HTTP_400_BAD_REQUEST,
        detail="Failed to decode text file. Ensure it is encoded in UTF-8 or standard plain text."
    )

def validate_image_bytes(file_bytes: bytes) -> None:
    """
    Uses PIL to verify that the file bytes form a valid, readable image.
    """
    try:
        img = Image.open(io.BytesIO(file_bytes))
        img.verify()
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Invalid or corrupt image file: {str(e)}"
        )

async def extract_content(file_bytes: bytes, filename: str) -> tuple[str | None, bytes | None, str]:
    """
    Route extraction based on validated file type.
    Returns:
        (extracted_text, image_bytes, file_type)
        For documents: (text, None, file_type)
        For images: (None, raw_bytes, file_type)
    """
    file_type = validate_file(filename, len(file_bytes))
    
    if file_type == "pdf":
        text = extract_pdf_text(file_bytes)
        return text, None, file_type
        
    elif file_type == "docx":
        text = extract_docx_text(file_bytes)
        return text, None, file_type
        
    elif file_type == "text":
        text = extract_txt_text(file_bytes)
        return text, None, file_type
        
    elif file_type == "image":
        validate_image_bytes(file_bytes)
        return None, file_bytes, file_type
        
    else:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Unhandled file type validation"
        )
